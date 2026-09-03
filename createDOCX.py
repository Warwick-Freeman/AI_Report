############################################
# The report as an editable Word document.
#
# The PDF is a fixed record; this is the same report in a form the
# electroencephalographer can change. SCORE's own report is a document a human
# signs, and a reader who cannot correct a wording, add a sentence about the
# clinical context or delete a finding they disagree with will either sign
# something they do not quite mean or retype the whole thing.
#
# Everything comes from the same results dict the PDF is built from, so the two
# cannot drift apart in substance. Real Word styles are used - Heading 1,
# Heading 2, Table Grid - rather than hand-set fonts, so the document restyles
# cleanly and takes a house template.
#
# The figures are the ones the analysis drew, embedded from the same files the
# PDF uses.
############################################
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from createPDF import FIGURE_CAPTIONS, figureNames
from recording import _hms

# Grey, for the basis lines that sit under a value. The same distinction the
# PDF makes between a finding and the measurement behind it.
BASIS_GREY = RGBColor(0x69, 0x69, 0x69)

# The usable width of a portrait page at these margins. Every column set below
# adds up to this or less; Word will otherwise widen the table past the margin.
PAGE_WIDTH_IN = 7.1

# The captions come from createPDF, so the two documents cannot label the same
# figure differently - which they did until the Word captions were corrected.


# Section tokens a template can place, in the order they are used when a
# template places none of them.
#
# The syntax is ProfusionEEG's own: [@SCORE,pdr] has the same shape as the
# [@99,ServiceDetails,GetTableValue,...] calls its report templates already
# carry, so a template is one artefact with one convention and anyone who has
# written a Profusion template already knows how to read it.
SECTION_ORDER = ('recording', 'pdr', 'interictal', 'episodes', 'events',
                 'sleep', 'artifacts', 'conclusion', 'narrative',
                 'measurements', 'figures')

# [@SCORE,pdr] / [@score , figures] / [@SCORE,all]
TOKEN_PATTERN = re.compile(r'\[@\s*SCORE\s*,\s*([A-Za-z_]+)\s*\]', re.I)


def writeDOCX(fileName, results, dest_folder, ai_report_text=None,
              template=None):
    """Write the report as a .docx. Returns the path written.

    Without a template the sections are written in their default order.

    With one, the template's own styles, header and footer carry - a logo in the
    Word header needs nothing more than that - and any [@SCORE,<section>] token
    in the body is replaced in place by that section, so a department decides
    where the report's parts sit among its own text. Tokens it does not use are
    appended at the end rather than dropped: leaving a template out of date must
    not silently remove a clinical section from a report.

    ProfusionEEG's own tokens - [@1000], [@99,...] - are left exactly as they
    are, for its pass to fill.
    """
    document = Document(template) if template else Document()
    if not template:
        _setUpStyles(document)

    excluded = set(results.get('_excluded') or ())
    stem = os.path.splitext(os.path.basename(fileName))[0]

    # Each section as something that can be rendered wherever it is wanted.
    renderers = {
        'recording': lambda d: _recording(d, results.get('recording')),
        'pdr': lambda d: _pdr(d, results.get('pdr')),
        'interictal': lambda d: _interictal(d, results.get('interictal')),
        'episodes': lambda d: _episodes(
            d, (results.get('spikeseizure') or {}).get('episodes'),
            (results.get('spikeseizure') or {}).get('notes')),
        'events': lambda d: _studyEvents(d, results.get('selected_events')),
        'sleep': lambda d: _sleep(d, results.get('sleep')),
        'artifacts': lambda d: _artifacts(d, results.get('artifacts')),
        'conclusion': lambda d: _conclusion(d, results.get('conclusion')),
        'narrative': lambda d: (_narrative(d, ai_report_text)
                                if ai_report_text else None),
        'measurements': lambda d: _measurements(d, results),
        'figures': lambda d: _figures(d, dest_folder, fileName),
    }
    # A section the reader excluded on the review screens is not rendered at
    # all, wherever a template puts it. Placement is the template's decision;
    # inclusion is the reader's, per study.
    for sectionId in list(renderers):
        if sectionId in excluded:
            renderers[sectionId] = lambda d: None

    placed, unknown = _placeTokens(document, renderers)

    if not placed:
        # No template, or a template that places nothing: the report in full,
        # in its own order, after whatever the template already held.
        document.add_heading('EEG Report - %s' % fileName, level=0)
        _scoreNote(document)
        for sectionId in SECTION_ORDER:
            renderers[sectionId](document)
    else:
        missing = [s for s in SECTION_ORDER if s not in placed]
        if missing:
            # Appended, not dropped. A template that has not caught up with a
            # new section would otherwise quietly leave it out of the report.
            print('Template placed %d section(s); appending %s, which it does '
                  'not position.' % (len(placed), ', '.join(missing)))
            document.add_heading('Further Findings', level=1)
            _basis(document, 'Sections the report template does not position. '
                             'Add a [@SCORE,<name>] token where each belongs.',
                   Inches(0))
            for sectionId in missing:
                renderers[sectionId](document)
        if unknown:
            print('Template has unrecognised section token(s): %s. Known names: '
                  '%s.' % (', '.join(sorted(unknown)), ', '.join(SECTION_ORDER)))

    if template:
        _templateNote(document, template)

    outFile = os.path.join(dest_folder, stem + '.docx')
    try:
        document.save(outFile)
    except PermissionError as e:
        # Word holds an exclusive lock on an open document, and regenerating a
        # report you are reading is an obvious thing to do. Say which file and
        # why rather than surfacing a bare errno.
        raise PermissionError(
            'Cannot write %s - it is open in another application, most likely '
            'Word. Close it and generate again.' % os.path.abspath(outFile)) from e
    outFile = os.path.abspath(outFile)
    print('Successfully generate docx file: ', outFile)
    return outFile


# --------------------------------------------------------------- placement

def _tokenParagraphs(document):
    """Body paragraphs that are a [@SCORE,...] token and nothing else.

    A token has to be alone in its paragraph. A section is a block, so a
    placeholder for one is a block too - and prose that merely mentions a token
    is not a placeholder. The sample template's own instructions say what
    [@SCORE,all] does, and matching tokens anywhere in a paragraph made that
    sentence render the whole report a second time. A department documenting
    its own template would hit exactly the same thing.

    Headers and footers are deliberately not searched: what belongs there is a
    logo and the patient tokens ProfusionEEG fills, not a findings table.

    Returns (placeholders, mentioned) - the second being paragraphs that
    contain a token among other text, so they can be reported rather than
    silently ignored.
    """
    placeholders, mentioned = [], []
    for paragraph in document.paragraphs:
        text = (paragraph.text or '').strip()
        if not text:
            continue
        match = TOKEN_PATTERN.fullmatch(text)
        if match:
            placeholders.append((paragraph, match.group(1).lower()))
        elif TOKEN_PATTERN.search(text):
            mentioned.append(text)
    return placeholders, mentioned


def _renderBefore(document, anchor, render):
    """Render a section and move what it produced to just before anchor.

    python-docx can only append, so the section is built at the end of the body
    and then moved. Identity is tracked rather than position because a renderer
    may add a section break as well as paragraphs and tables.
    """
    body = document.element.body
    # The list is held, not a set of id()s. lxml builds element proxies on
    # demand and only guarantees the same proxy while a reference is alive, so
    # an id() taken from a throwaway proxy can be reused by a different node -
    # which made this treat paragraphs already in the template as newly added
    # and move them, scrambling the document.
    before = list(body)
    render(document)
    added = [e for e in body if not any(e is kept for kept in before)]
    for element in added:
        anchor.addprevious(element)


def _removeToken(paragraph):
    """Drop the placeholder paragraph now its section has been rendered."""
    element = paragraph._p
    element.getparent().remove(element)


def _placeTokens(document, renderers):
    """Fill every [@SCORE,...] token in the template. Returns (placed, unknown).

    Only these tokens are touched. ProfusionEEG's own - [@1000], [@99,...] -
    are left byte for byte as they are, so its pass can fill them whether it
    runs before or after this one.
    """
    tokens, mentioned = _tokenParagraphs(document)
    for text in mentioned:
        # Told about, not acted on: a token sharing a paragraph with other text
        # is far more likely to be someone writing about it than placing it.
        print('Ignored a section token inside other text (a token must be alone '
              'in its paragraph): %r' % text[:70])
    if not tokens:
        return set(), set()

    placed, unknown = set(), set()
    for paragraph, sectionId in tokens:
        if sectionId == 'all':
            for name in SECTION_ORDER:
                if name in placed:
                    continue
                _renderBefore(document, paragraph._p, renderers[name])
                placed.add(name)
            _removeToken(paragraph)
            continue
        if sectionId not in renderers:
            unknown.add(sectionId)
            continue
        if sectionId in placed:
            # A template naming the same section twice would otherwise report
            # the same findings twice.
            print('Section token [@SCORE,%s] appears more than once; the later '
                  'one was ignored.' % sectionId)
            _removeToken(paragraph)
            continue
        _renderBefore(document, paragraph._p, renderers[sectionId])
        placed.add(sectionId)
        _removeToken(paragraph)
    return placed, unknown


def _templateNote(document, template):
    """Which template produced this report.

    A department edits its template over time, and 'which letterhead and
    boilerplate was this signed under' is the question asked afterwards.
    """
    _basis(document,
           'Report template: %s' % os.path.basename(template), Inches(0))


# ------------------------------------------------------------------ scaffolding

def _setUpStyles(document):
    """Body text a clinician can read, and margins a table fits in."""
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    for section in document.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def _scoreNote(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        'Scored against SCORE - Standardized Computer-based Organized Reporting '
        'of EEG, second version (Beniczky et al., Clin Neurophysiol '
        '2017;128:2334-2346). Values marked as not scored were not determinable '
        'from the recording; they are not negative findings.')
    run.font.size = Pt(8)
    run.font.color.rgb = BASIS_GREY


def _basis(document, text, indent=Inches(0.25)):
    """The measurement under a value, in the smaller grey the PDF uses."""
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = indent
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = BASIS_GREY


def _repeatHeader(row):
    """Mark a row as a header so Word repeats it on every page.

    A table that runs over a page break otherwise continues with unlabelled
    columns, which for a findings table means a reader cannot tell prevalence
    from confidence.
    """
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    properties.append(header)


def _table(document, headers, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Word recalculates column widths when autofit is left on, so the widths set
    # per cell below were being discarded and the columns came out squashed.
    table.autofit = False
    header = table.rows[0].cells
    for index, text in enumerate(headers):
        header[index].text = ''
        run = header[index].paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        if widths:
            header[index].width = Inches(widths[index])
    _repeatHeader(table.rows[0])
    return table


def _row(table, values, widths=None, basis=None):
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = ''
        run = cells[index].paragraphs[0].add_run(
            '' if value is None else str(value))
        run.font.size = Pt(9)
        if widths:
            cells[index].width = Inches(widths[index])
    if basis:
        paragraph = cells[0].add_paragraph()
        run = paragraph.add_run(basis)
        run.font.size = Pt(7.5)
        run.font.color.rgb = BASIS_GREY
    return cells


def _notes(document, notes, heading='Notes'):
    notes = [n for n in (notes or []) if n]
    if not notes:
        return
    document.add_heading(heading, level=2)
    for note in notes:
        document.add_paragraph(note, style='List Bullet')


def _readerFields(document, fields, entries=None):
    """What SCORE wants from the reader.

    An answer given before the document was generated is printed. Only what was
    left unanswered stays a blank line - a field the reader filled in and then
    saw printed as underscores would send them back to type it a second time.
    """
    entries = entries or {}
    answered = [f for f in fields if entries.get(f)]
    unanswered = [f for f in fields if not entries.get(f)]

    if answered:
        for field in answered:
            paragraph = document.add_paragraph(style='List Bullet')
            paragraph.add_run('%s: ' % field).bold = True
            paragraph.add_run(str(entries[field]))
    if unanswered:
        document.add_heading('For completion by the reader', level=2)
        for field in unanswered:
            paragraph = document.add_paragraph(style='List Bullet')
            paragraph.add_run('%s: ' % field).bold = True
            paragraph.add_run('_' * 40)


# --------------------------------------------------------------- the sections

def _recording(document, recording):
    if not recording:
        return
    document.add_heading('Patient and Recording Conditions', level=1)

    for group, title in (('patient', 'Patient'), ('conditions', 'Recording')):
        entries = (recording.get(group) or {})
        if not entries:
            continue
        document.add_heading(title, level=2)
        table = _table(document, ('Item', 'Value'), (2.0, 5.0))
        for label, value in entries.items():
            _row(table, (label, value), (2.0, 5.0))

    lines = recording.get('duration_lines') or []
    if lines:
        document.add_heading('Duration examined', level=2)
        for line in lines:
            document.add_paragraph(line, style='List Bullet')

    _activation(document, recording.get('activation'))

    fields = recording.get('technologist_fields') or []
    if fields:
        _readerFields(document, fields,
                      recording.get('technologist_entries'))

    _notes(document, recording.get('notes'))


def _activation(document, activation):
    procedures = (activation or {}).get('procedures') or []
    if not procedures:
        return
    document.add_heading('Activation procedures', level=2)
    widths = (1.7, 1.6, 2.3, 1.4)
    table = _table(document, ('Procedure', 'State', 'Timing', 'Response'), widths)
    for procedure in procedures:
        onset = procedure.get('onset_seconds')
        _row(table,
             (procedure.get('name'), procedure.get('state'),
              '' if onset is None else 'from %s' % _hms(onset),
              procedure.get('response') or 'Not scored'),
             widths, basis=procedure.get('detail'))
    for note in activation.get('notes') or []:
        _basis(document, note, indent=Inches(0))


def _pdr(document, pdr):
    if not pdr:
        return
    document.add_heading('Posterior Dominant Rhythm', level=1)
    widths = (1.6, 2.2, 0.9, 2.3)
    table = _table(document, ('Property', 'Scored value', 'Confidence',
                              'Measurement behind it'), widths)

    labels = (('frequency', 'Frequency'), ('frequency_asymmetry', 'Frequency asymmetry'),
              ('amplitude', 'Amplitude'), ('amplitude_asymmetry', 'Amplitude asymmetry'),
              ('reactivity', 'Reactivity'), ('organization', 'Organisation'),
              ('caveat', 'Caveat'), ('absence', 'Absence of the PDR'),
              ('significance', 'Significance'))
    overridden = []
    for key, label in labels:
        entry = pdr.get(key)
        if not isinstance(entry, dict):
            continue
        term = entry.get('term')
        if entry.get('provisional'):
            term = '%s (provisional)' % term
        _row(table, (label, term, entry.get('confidence', ''),
                     entry.get('basis', '')), widths)
        if entry.get('overridden'):
            overridden.append((label, entry))

    if overridden:
        document.add_heading('Overridden by the reader', level=2)
        for label, entry in overridden:
            document.add_paragraph(
                '%s: reported as %s; measured %s'
                % (label, entry.get('term'), entry.get('measured_term')),
                style='List Bullet')

    _notes(document, pdr.get('notes'))


def _findingTable(document, findings, timingHeader):
    widths = (1.5, 2.0, 1.4, 1.3, 0.8)
    table = _table(document, ('Graphoelement', 'Location', timingHeader,
                              'Mode of appearance', 'Confidence'), widths)
    for finding in findings:
        location = (finding.get('location') or {})
        timing = finding.get('prevalence') or finding.get('incidence') or ''
        if finding.get('count'):
            timing = ('%s (%d)' % (timing, finding['count'])) if timing \
                else str(finding['count'])
        mode = ', '.join(x for x in (finding.get('mode_of_appearance'),
                                     finding.get('discharge_pattern')) if x)
        basis = '. '.join(x for x in (finding.get('basis'),
                                      finding.get('timing_basis')) if x)
        _row(table,
             (finding.get('name'), location.get('text') or '', timing, mode,
              finding.get('confidence', '')), widths, basis=basis)


def _interictal(document, interictal):
    findings = (interictal or {}).get('findings') or []
    if not interictal:
        return
    document.add_heading('Interictal Findings', level=1)
    if findings:
        _findingTable(document, findings, 'Prevalence / incidence')
    else:
        document.add_paragraph(
            'Analysed. No abnormal interictal rhythmic activity or epileptiform '
            'discharges of this class were detected. This is a clinical result, '
            'not a failure to analyse.')
    _notes(document, interictal.get('notes'),
           heading='What the analysis decided not to report')


def _episodes(document, episodes, notes):
    if not episodes:
        return
    document.add_heading('Episodes', level=1)
    widths = (1.9, 1.9, 1.0, 1.5, 0.8)
    table = _table(document, ('Episode', 'Location', 'Onset', 'Duration',
                              'Confidence'), widths)
    readerFields = None
    for episode in episodes:
        location = (episode.get('location') or {})
        onset = episode.get('onset_seconds')
        duration = episode.get('duration_seconds')
        _row(table,
             (episode.get('name'), location.get('text') or '',
              '' if onset is None else _hms(onset),
              episode.get('duration_band')
              or ('' if duration is None else '%.0f s' % duration),
              episode.get('confidence', '')),
             widths, basis=episode.get('basis'))
        readerFields = episode.get('reader_fields') or readerFields

    # Each episode's own answers, under its own heading: with more than one
    # seizure, a single pooled list could not say which episode it described.
    for index, episode in enumerate(episodes):
        entries = episode.get('reader_entries') or {}
        onset = episode.get('onset_seconds')
        where = ('the episode at %s' % _hms(onset)) if onset is not None \
            else 'episode %d' % (index + 1)
        document.add_heading('Clinical detail - %s' % where, level=2)
        _readerFields(document,
                      episode.get('reader_fields') or list(entries),
                      entries)
    _notes(document, notes)


def _studyEvents(document, events):
    if not events:
        return
    document.add_heading('Study Events', level=1)
    _basis(document, 'Selected by the reader as context, from the study\'s own '
                     'event record. None of these is a finding.', Inches(0))
    widths = (1.7, 0.9, 0.7, 2.1, 1.6)
    table = _table(document, ('Event', 'Time', 'Duration', 'Text', 'Channels'),
                   widths)
    for event in events:
        seconds = event.get('seconds')
        duration = event.get('duration_seconds') or 0
        note = []
        if event.get('provocation'):
            note.append('SCORE provocation: %s' % event['provocation'])
        if event.get('is_detection'):
            note.append('detector output')
        _row(table,
             (event.get('type'), '' if seconds is None else _hms(seconds),
              '%.0f s' % duration if duration else '',
              event.get('text') or '',
              ', '.join(event.get('channels') or [])),
             widths, basis='. '.join(note) or None)


def _sleep(document, sleep):
    if not sleep:
        return
    document.add_heading('Sleep and Drowsiness', level=1)
    backend = sleep.get('backend') or 'model'
    _basis(document, 'Staged by %s. This is AI analysis and has not been '
                     'verified.' % backend, Inches(0))

    rows = [(label, sleep.get(key)) for label, key in
            (('Stages reached', 'stages_text'), ('Sleep reached', 'sleep_reached'),
             ('Drowsiness', 'drowsiness'))
            if sleep.get(key) is not None]
    if rows:
        table = _table(document, ('Item', 'Value'), (2.0, 5.0))
        for label, value in rows:
            _row(table, (label, value), (2.0, 5.0))

    findings = sleep.get('findings') or []
    if findings:
        document.add_heading('Sleep graphoelements', level=2)
        _findingTable(document, findings, 'Prevalence / incidence')
    _notes(document, sleep.get('notes'))


def _artifacts(document, artifacts):
    findings = (artifacts or {}).get('findings') or []
    if not artifacts:
        return
    document.add_heading('EEG Artifacts', level=1)
    if findings:
        widths = (1.8, 2.2, 1.6, 0.9)
        table = _table(document, ('Artifact type', 'Location',
                                  'Prevalence / incidence', 'Confidence'), widths)
        for finding in findings:
            location = (finding.get('location') or {})
            timing = finding.get('prevalence') or finding.get('incidence') or ''
            _row(table,
                 (finding.get('name'), location.get('text') or '', timing,
                  finding.get('confidence', '')), widths,
                 basis=finding.get('basis'))
        judged = {f['name']: f['significance'] for f in findings
                  if f.get('significance')}
        if judged:
            document.add_heading('Significance', level=2)
            for name, significance in judged.items():
                paragraph = document.add_paragraph(style='List Bullet')
                paragraph.add_run('%s: ' % name).bold = True
                paragraph.add_run(significance)
        unjudged = [f['name'] for f in findings if not f.get('significance')]
        if unjudged:
            _readerFields(document, ['Significance - %s' % n for n in unjudged])
    else:
        document.add_paragraph('Analysed. No artifact types were detected.')
    _notes(document, artifacts.get('notes'))


def _conclusion(document, conclusion):
    document.add_heading('Diagnostic Significance and Conclusion', level=1)
    _basis(document, 'SCORE reserves the diagnostic significance and the '
                     'conclusion for the electroencephalographer. Nothing in '
                     'the analysis writes them.', Inches(0))

    conclusion = conclusion or {}
    table = _table(document, ('Item', 'Value'), (2.0, 5.0))
    for label, key in (('Diagnostic significance', 'significance'),
                       ('Diagnostic yield', 'yield')):
        _row(table, (label, conclusion.get(key) or 'Not scored'), (2.0, 5.0))
    for key in ('basis', 'notes'):
        for line in (conclusion.get(key) or []):
            _basis(document, line)

    summary = conclusion.get('summary') or conclusion.get('text')
    if summary:
        document.add_heading('Summary and clinical comments', level=2)
        document.add_paragraph(summary)

    # Signed by, and when. Printed where given, asked for where not.
    signature = {'Reported by': conclusion.get('reported_by'),
                 'Date': conclusion.get('report_date')}
    document.add_heading('Sign-off', level=2)
    _readerFields(document, list(signature), signature)

    outstanding = [label for label, key in
                   (('Diagnostic significance', 'significance'),
                    ('Diagnostic yield', 'yield'),
                    ('Summary and clinical comments', 'summary'))
                   if not conclusion.get(key)]
    if outstanding:
        _readerFields(document, outstanding)


def _narrative(document, text):
    document.add_heading('Narrative Draft', level=1)
    _basis(document, 'Drafted by a language model. This is AI analysis, has not '
                     'been verified, and is offered as a draft rather than as a '
                     'scored value - edit or delete it.', Inches(0))
    for block in str(text).split('\n'):
        if block.strip():
            document.add_paragraph(block.strip())


def _measurements(document, results):
    """The underlying numbers, for anyone who wants to check the findings."""
    document.add_heading('Underlying Measurements', level=1)
    widths = (2.6, 1.4, 1.4, 1.4)
    table = _table(document, ('Measurement', 'Left', 'Right', 'Total'), widths)

    def number(value, unit=''):
        if value is None:
            return ''
        try:
            return '%.2f%s' % (float(value), unit)
        except (TypeError, ValueError):
            return str(value)

    rows = (
        ('Background frequency (Hz)', 'left_backgroud_frequency',
         'right_backgroud_frequency', None),
        ('Slow-wave ratio', 'left_slow_ratio', 'right_slow_ratio',
         'total_slow_ratio'),
        ('Beta ratio', 'left_beta_ratio', 'right_beta_ratio', 'total_beta_ratio'),
        ('Anterior-posterior difference', 'left_AP_difference',
         'right_AP_difference', 'AP_difference'),
    )
    for label, left, right, total in rows:
        _row(table, (label, number(results.get(left)), number(results.get(right)),
                     number(results.get(total)) if total else ''), widths)

    bad = results.get('bad_channels')
    _basis(document, 'Electrodes marked bad: %s'
           % (', '.join(bad) if bad else 'none'), Inches(0))
    ratio = results.get('removeEpochsRatio')
    if ratio is not None:
        _basis(document, 'Epochs rejected as artifact: %s'
               % number(float(ratio) * 100, '%'), Inches(0))


def _usableBox(section):
    """The space a figure has on the page, in inches."""
    width = (section.page_width.inches
             - section.left_margin.inches - section.right_margin.inches)
    height = (section.page_height.inches
              - section.top_margin.inches - section.bottom_margin.inches)
    return width, height


def _figureWidth(path, maxWidth, maxHeight):
    """How wide to place a figure so it fits the page in both directions.

    Fitting width alone is what made the figures section unusable: the
    topographic maps are 1440x1800 and 1440x2160, so at 9.5 inches wide they
    wanted 11.9 and 14.2 inches of height on a page with 6.5 - Word moved each
    one to the page after its heading and left the heading stranded.
    """
    try:
        from PIL import Image
        with Image.open(path) as image:
            pixelWidth, pixelHeight = image.size
        aspect = float(pixelWidth) / float(pixelHeight or 1)
    except Exception:
        return maxWidth
    if aspect <= 0:
        return maxWidth
    # Whichever bound binds first.
    return min(maxWidth, maxHeight * aspect)


def _figures(document, dest_folder, fileName):
    """The analysis figures, each below its own caption.

    The page stays portrait. Four of the six figures are portrait or square -
    the topographic maps and the spectrogram - so a landscape section made them
    smaller, not larger, as well as breaking the document into two orientations
    for anyone editing it.
    """
    names = figureNames(fileName)
    # Paired with the index that names them: if one figure is missing, the rest
    # must keep their own captions rather than shifting up by one.
    present = [(index, os.path.join(dest_folder, name))
               for index, name in enumerate(names)]
    present = [(i, p) for i, p in present if os.path.isfile(p)]
    if not present:
        return

    document.add_page_break()
    document.add_heading('Figures', level=1)

    maxWidth, maxHeight = _usableBox(document.sections[-1])
    # Room for the caption above the figure.
    maxHeight -= 0.6

    for index, path in present:
        caption = FIGURE_CAPTIONS.get(index) or os.path.basename(path)
        heading = document.add_heading(caption, level=2)
        # Word must not put the caption on one page and the figure on the next.
        heading.paragraph_format.keep_with_next = True

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        paragraph.add_run().add_picture(
            path, width=Inches(_figureWidth(path, maxWidth, maxHeight)))
