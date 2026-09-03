############################################
# Build the sample report template.
#
# Run this to regenerate templates/Sample EEG Report.docx. It is a starting
# point for a department: replace the logo placeholder with your own image,
# change the wording, move the [@SCORE,...] tokens to where you want each part
# of the report, and delete the guidance page.
#
#   python make_sample_template.py
#
# Two token vocabularies live side by side in it, deliberately:
#
#   [@1000] and [@99,...]   ProfusionEEG's own, filled by its report pass.
#                           Nothing here touches them.
#   [@SCORE,pdr] and so on  filled by this project's Word writer.
#
# They share one syntax so a template is one artefact with one convention.
############################################
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import createDOCX

GREY = RGBColor(0x69, 0x69, 0x69)
OUT = os.path.join('templates', 'Sample EEG Report.docx')


def guidance(document, text, size=8.5):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GREY
    run.italic = True
    return paragraph


def token(document, name, label):
    """A section token with a line saying what it will become."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run('[@SCORE,%s]' % name)
    run.bold = True
    guidance(document, '   ^ replaced by: %s' % label)


def build():
    document = Document()

    # Header: where a hospital logo goes. A picture in the Word header repeats
    # on every page and needs nothing from this project.
    header = document.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('[ replace with your hospital logo: Insert > Pictures, '
                         'in this header ]')
    run.font.size = Pt(9)
    run.font.color.rgb = GREY

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Confidential patient record')
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

    document.add_heading('Electroencephalogram Report', level=0)

    # ProfusionEEG's own tokens, exactly as its shipped templates use them.
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    for label, value in (
            ('Surname', '[@1000]'),
            ('Given name(s)', '[@1001]'),
            ('UR number', '[@1002]'),
            ('Date of birth', '[@1004]'),
            ('Date of test', '[@2000]'),
            ('Referred by', '[@99,ServiceDetails,GetReferringPhysicianValue,'
                            'firstname] [@99,ServiceDetails,'
                            'GetReferringPhysicianValue,lastname]'),
            ('Reported by', '[@99,ServiceDetails,GetStaffResourceValue,'
                            'Reporting Neurologist,firstName] '
                            '[@99,ServiceDetails,GetStaffResourceValue,'
                            'Reporting Neurologist,lastName]')):
        cells = table.add_row().cells
        cells[0].width = Inches(1.8)
        cells[1].width = Inches(5.2)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(value)
    guidance(document, 'The tokens above are ProfusionEEG\'s and are filled by '
                       'its report pass. This project does not touch them.')

    # A department's own section, before any findings.
    document.add_heading('Clinical question and referral', level=1)
    document.add_paragraph('_' * 70)
    document.add_paragraph('_' * 70)
    guidance(document, 'Your own text. Keep, change or delete it - the report '
                       'is built around whatever is here.')

    document.add_heading('Technical description', level=1)
    token(document, 'recording', 'patient and recording conditions, the '
                                 'duration examined, and the activation '
                                 'procedures performed')

    document.add_heading('Findings', level=1)
    token(document, 'pdr', 'the nine SCORE posterior dominant rhythm properties')
    token(document, 'interictal', 'interictal findings, including any spike '
                                  'detections')
    token(document, 'episodes', 'detected electrographic seizures')
    token(document, 'sleep', 'sleep staging and sleep graphoelements')
    token(document, 'artifacts', 'artifact types, with their coverage')

    document.add_heading('Departmental protocol notes', level=1)
    document.add_paragraph('_' * 70)
    guidance(document, 'Another of your own sections, sitting between report '
                       'parts. This is what the tokens are for.')

    document.add_heading('Interpretation', level=1)
    token(document, 'conclusion', 'diagnostic significance and the conclusion, '
                                  'with the fields SCORE leaves to the reader')
    token(document, 'narrative', 'the language-model narrative draft, if one '
                                 'was requested')

    document.add_heading('Appendix', level=1)
    token(document, 'events', 'the study events selected during review')
    token(document, 'measurements', 'the underlying numbers')
    token(document, 'figures', 'the analysis figures')

    document.add_page_break()
    document.add_heading('How to use this template', level=1)
    for line in (
        'Save a copy under your own name in this templates folder, then choose '
        'it in Settings.',
        'Put your logo in the Word header, not in the body: a header repeats on '
        'every page.',
        'Move any [@SCORE,<name>] token to where you want that part of the '
        'report. Add your own headings and text freely around them.',
        'A token you delete does not remove the section. It is appended at the '
        'end under "Further Findings" instead, because a template that has not '
        'caught up must not quietly drop a clinical section. To leave a section '
        'out of a particular report, exclude it on the review screens.',
        '[@SCORE,all] places the whole report at that point, for a template '
        'that only wants a letterhead.',
        'Section names: ' + ', '.join(createDOCX.SECTION_ORDER) + '.',
        'Styles used by the report: Heading 1, Heading 2, Normal, Table Grid, '
        'List Bullet. Restyle those and the report follows.',
        'Delete this page before using the template.',
    ):
        document.add_paragraph(line, style='List Bullet')

    os.makedirs('templates', exist_ok=True)
    document.save(OUT)
    print('Wrote %s' % os.path.abspath(OUT))
    return OUT


if __name__ == '__main__':
    build()
